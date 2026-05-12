from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "jour2_organisation" / "Jour_2_Organisation_Ping_Pang.docx"


ACCENT = RGBColor(17, 94, 89)
DARK = RGBColor(32, 40, 45)
MUTED = RGBColor(91, 103, 112)
LIGHT = "E8F3F1"
HEADER = "DDEDEA"
GRID = "BFCBC8"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def borders(cell, color=GRID):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        node = tc_borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_cell(cell, bold=False, fill=None, color=None, size=9.5):
    if fill:
        shade(cell, fill)
    borders(cell)
    margins(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(size)
            run.font.bold = bold
            if color:
                run.font.color.rgb = color


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.line_spacing = 1.08
    styles["Normal"].paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 16, ACCENT),
        ("Heading 2", 13, DARK),
        ("Heading 3", 11, DARK),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.text = "Ping Pang x Eugenia — Livrable Jour 2"
    header.runs[0].font.name = "Arial"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    for run in footer.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = MUTED
    return doc


def add_title(doc):
    p = doc.add_paragraph()
    p.style = doc.styles["Title"]
    r = p.add_run("Jour 2 — Organisation du groupe")
    r.font.name = "Arial"
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = ACCENT
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.style = doc.styles["Subtitle"]
    r = p.add_run("Répartition des rôles, roadmap, objectifs journaliers et gestion des risques")
    r.font.name = "Arial"
    r.font.size = Pt(11)
    r.font.color.rgb = MUTED
    p.paragraph_format.space_after = Pt(12)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, LIGHT)
    borders(cell, "9FB8B4")
    margins(cell, top=150, bottom=150, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.font.name = "Arial"
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = ACCENT
    p = cell.add_paragraph()
    r = p.add_run(body)
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r.font.color.rgb = DARK


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        style_cell(cell, bold=True, fill=HEADER, color=ACCENT, size=9)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            style_cell(cells[i], size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(5)
        p.add_run(item)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(5)
        p.add_run(item)


def build():
    doc = setup_doc()
    add_title(doc)

    add_callout(
        doc,
        "Rappel de la règle d'organisation",
        "Le groupe est divisé en deux équipes produit. Chaque membre reste sur son produit. La seule personne autorisée à travailler sur les deux produits est le chef de projet, afin d'assurer la cohérence globale.",
    )

    doc.add_heading("1. Objectif du livrable Jour 2", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "Ce document présente l'organisation du groupe pour la suite du projet Ping Pang. Il définit le chef de projet, les deux équipes produit, la roadmap, les objectifs journaliers, les responsabilités de chacun et les principaux risques à surveiller."
    )

    doc.add_heading("2. Objectifs finaux du projet", level=1)
    add_table(
        doc,
        ["Produit", "Objectif", "Résultat attendu"],
        [
            [
                "Player Profile / Training App",
                "Construire le profil digital du joueur et lui permettre de suivre sa progression.",
                "Un espace joueur clair avec données de profil, historique, statistiques, entraînement et analyse de performance.",
            ],
            [
                "Ranking Matchmaking",
                "Permettre aux joueurs de trouver des adversaires, lancer des matchs, enregistrer les scores et alimenter un ranking.",
                "Un parcours simple pour matcher, jouer, scorer, valider le résultat et faire évoluer le classement.",
            ],
        ],
        [2300, 3400, 3660],
    )

    doc.add_heading("3. Organisation de l'équipe", level=1)
    add_table(
        doc,
        ["Rôle", "Personne(s)", "Périmètre"],
        [
            [
                "Chef de projet / PM",
                "Clément",
                "Coordination globale, roadmap, cohérence produit, scraping/data, système Elo/PGR, gestion des risques et lien entre les deux apps.",
            ],
            [
                "Équipe Ranking Matchmaking",
                "Corentin, Agathe",
                "Recherche de joueurs, matchmaking, lancement de match, scoring, validation du résultat, classement et historique côté ranking.",
            ],
            [
                "Équipe Player Profile / Training App",
                "Gaspard, Tara",
                "Profil joueur, suivi d'entraînement, récupération des données de match, statistiques, analyse de performance et coach IA.",
            ],
        ],
        [2200, 2200, 4960],
    )

    doc.add_heading("4. Rôle du chef de projet", level=1)
    p = doc.add_paragraph()
    p.add_run("Le chef de projet est le seul membre transversal. Son rôle est de garder une vision globale sans remplacer les équipes produit.")
    bullets(
        doc,
        [
            "Coordonner les deux équipes et organiser les points d'avancement.",
            "Maintenir la roadmap et vérifier que chaque journée produit un livrable concret.",
            "Garantir la cohérence entre Ranking Matchmaking et Training App.",
            "Piloter les sujets transversaux : scraping, data, modèle Elo/PGR et architecture commune des matchs.",
            "S'assurer que les deux produits respectent la vibe Ping Pang : simple, utile, élégant, sportif et premium.",
            "Identifier les risques, arbitrer les priorités et débloquer les points bloquants.",
        ],
    )

    doc.add_heading("5. Responsabilités par équipe produit", level=1)
    doc.add_heading("5.1 Équipe Ranking Matchmaking — Corentin et Agathe", level=2)
    bullets(
        doc,
        [
            "Définir le parcours utilisateur : trouver un joueur, lancer un match, scorer, valider le résultat.",
            "Concevoir l'interface de matchmaking et de challenge.",
            "Construire la logique de match : joueurs, statut, score, sets, vainqueur, validation.",
            "Préparer l'affichage du classement et de l'historique des matchs.",
            "Collaborer avec le PM uniquement pour les données communes et le système Elo/PGR.",
        ],
    )

    doc.add_heading("5.2 Équipe Player Profile / Training App — Gaspard et Tara", level=2)
    bullets(
        doc,
        [
            "Définir le profil joueur : informations, niveau, club, style de jeu, matériel et historique.",
            "Concevoir l'expérience de suivi d'entraînement et de progression.",
            "Prévoir l'intégration automatique des matchs créés depuis Ranking Matchmaking.",
            "Construire les statistiques utiles au joueur : progression, régularité, résultats, points forts/faibles.",
            "Réfléchir au coach IA : quelles données il reçoit et quelles recommandations il peut produire.",
        ],
    )

    doc.add_heading("6. Point clé : donnée de match commune", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "Les deux apps sont séparées côté produit, mais elles doivent partager une même donnée de match. L'utilisateur ne doit pas avoir à lancer ou saisir deux fois le même match."
    )
    add_table(
        doc,
        ["Étape", "Ranking Matchmaking", "Training App"],
        [
            ["1", "Crée le match et choisit les joueurs.", "Attend la donnée commune."],
            ["2", "Gère le scoring point par point ou set par set.", "Récupère les événements du match."],
            ["3", "Valide le résultat et met à jour le ranking.", "Analyse la performance et alimente le profil joueur."],
            ["4", "Affiche l'impact sur le classement.", "Affiche les statistiques et recommandations d'entraînement."],
        ],
        [900, 4200, 4260],
    )

    doc.add_heading("7. Roadmap de travail", level=1)
    add_table(
        doc,
        ["Période", "Objectif", "Livrables attendus"],
        [
            [
                "Jour 2",
                "Organisation du groupe et clarification des responsabilités.",
                "Équipes définies, roadmap, objectifs journaliers, risques et règles de travail.",
            ],
            [
                "Jours 3-4",
                "Cadrage produit et parcours utilisateurs.",
                "User flows, priorisation des fonctionnalités et premières maquettes fonctionnelles.",
            ],
            [
                "Semaine 1",
                "Spécifications et architecture.",
                "Modèle de données, choix techniques, structure des deux produits et données communes.",
            ],
            [
                "Semaine 2",
                "Construction du coeur produit.",
                "Profil joueur, création match, scoring, premières pages fonctionnelles.",
            ],
            [
                "Semaine 3",
                "Produit utilisable et intégration.",
                "Connexion Ranking/Training, statistiques, historique, ranking et parcours fluides.",
            ],
            [
                "Semaine 4",
                "Tests, polish et préparation démo.",
                "Produit stable, corrections UX, documentation technique et démo live.",
            ],
        ],
        [1600, 3300, 4460],
    )

    doc.add_heading("8. Objectifs journaliers", level=1)
    add_table(
        doc,
        ["Personne / équipe", "Objectifs journaliers"],
        [
            [
                "Clément — PM",
                "Faire un point matin/soir, suivre la roadmap, gérer les risques, arbitrer les priorités, cadrer scraping + Elo/PGR, vérifier la cohérence entre les deux produits.",
            ],
            [
                "Corentin",
                "Avancer sur le parcours Ranking : recherche joueur, challenge, création match ou scoring selon la priorité du jour.",
            ],
            [
                "Agathe",
                "Avancer sur l'UX/UI Ranking : matchmaking, validation du score, classement, historique et fluidité du parcours.",
            ],
            [
                "Gaspard",
                "Avancer sur la structure Training : profil joueur, données de progression, historique et intégration des matchs.",
            ],
            [
                "Tara",
                "Avancer sur l'UX/UI Training : dashboard joueur, stats, suivi d'entraînement et premières pistes coach IA.",
            ],
        ],
        [2200, 7160],
    )

    doc.add_heading("9. Gestion des risques", level=1)
    add_table(
        doc,
        ["Risque", "Impact", "Prévention"],
        [
            [
                "Trop de fonctionnalités",
                "Produit incomplet ou difficile à démontrer.",
                "Prioriser les parcours essentiels et valider chaque ajout avec le PM.",
            ],
            [
                "Deux apps qui ne communiquent pas",
                "Double saisie utilisateur et perte de valeur pour Training.",
                "Créer une donnée de match commune dès le départ.",
            ],
            [
                "Système Elo/PGR trop complexe",
                "Blocage technique et perte de temps.",
                "Commencer par une version simple, documentée, puis améliorer.",
            ],
            [
                "Scraping difficile ou accès API bloqué",
                "Manque de données réelles.",
                "Prévoir données mockées propres + demande API FFTT + sources publiques prioritaires.",
            ],
            [
                "Manque de cohérence visuelle Ping Pang",
                "Produit moins crédible face au jury.",
                "Suivre les brand guidelines et valider régulièrement la direction design.",
            ],
            [
                "Membres dispersés entre les deux produits",
                "Responsabilités floues et perte d'efficacité.",
                "Respecter la règle : seul le PM travaille sur les deux produits.",
            ],
        ],
        [2600, 3000, 3760],
    )

    doc.add_heading("10. Règles de fonctionnement", level=1)
    numbered(
        doc,
        [
            "Chaque membre reste sur son produit attribué.",
            "Le PM est le seul point de coordination entre les deux équipes.",
            "Chaque journée doit produire un livrable visible ou vérifiable.",
            "Les décisions produit sont documentées pour pouvoir être justifiées.",
            "Les deux produits doivent rester simples, utiles et cohérents avec l'univers Ping Pang.",
            "Les problèmes bloquants sont remontés rapidement au PM.",
        ],
    )

    doc.add_heading("11. Synthèse", level=1)
    add_callout(
        doc,
        "Organisation retenue",
        "Le groupe est structuré en deux pôles produit distincts : Corentin et Agathe sur Ranking Matchmaking, Gaspard et Tara sur Player Profile / Training App. Clément assure le rôle de chef de projet transversal, avec la responsabilité de la roadmap, du scraping, du système Elo/PGR, de la cohérence produit et du lien entre les deux applications.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
