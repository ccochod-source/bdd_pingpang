from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "pgr_doc" / "Brief_Ping_Pang_Global_Rating.docx"


ACCENT = RGBColor(17, 94, 89)
MUTED = RGBColor(91, 103, 112)
LIGHT = "E8F3F1"
TABLE_HEADER = "DDEDEA"
GRID = "BFCBC8"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=GRID, size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_cell(cell, bold=False, fill=None, color=None, size=10):
    if fill:
        set_cell_shading(cell, fill)
    set_cell_border(cell)
    set_cell_margins(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(size)
            run.font.bold = bold
            if color:
                run.font.color.rgb = color


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.style = doc.styles["Title"]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.font.color.rgb = ACCENT
    run.font.bold = True
    run.font.size = Pt(22)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.style = doc.styles["Subtitle"]
    r = p.add_run(subtitle)
    r.font.color.rgb = MUTED
    r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(12)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT)
    set_cell_border(cell, "9FB8B4")
    set_cell_margins(cell, top=160, bottom=160, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(11)
    r.font.color.rgb = ACCENT
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(body)
    r.font.name = "Arial"
    r.font.size = Pt(10)
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(5)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(5)
        p.add_run(item)


def add_small_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        style_cell(cell, bold=True, fill=TABLE_HEADER, color=ACCENT, size=9)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            style_cell(cells[i], size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.line_spacing = 1.08
    styles["Normal"].paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 16, ACCENT),
        ("Heading 2", 13, RGBColor(35, 48, 55)),
        ("Heading 3", 11, RGBColor(35, 48, 55)),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.text = "Ping Pang Global Rating — Brief produit & data"
    header.runs[0].font.name = "Arial"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    for run in footer.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = MUTED
    return doc


def build():
    doc = setup_doc()

    add_title(
        doc,
        "Ping Pang Global Rating (PGR)",
        "Document de cadrage : comment créer un rating mondial unique, avec un focus précis sur les joueurs professionnels.",
    )

    add_callout(
        doc,
        "Idée centrale",
        "Ping Pang ne fusionne pas directement les classements nationaux. Ping Pang crée un rating global indépendant, initialisé à partir des sources officielles, puis stabilisé par les matchs réels.",
    )

    doc.add_heading("1. Objectif du PGR", level=1)
    p = doc.add_paragraph()
    p.add_run("Le PGR, Ping Pang Global Rating, est le score commun qui permet de comparer tous les joueurs dans un même classement : débutants, amateurs, licenciés, compétiteurs nationaux et joueurs pros.")
    add_bullets(
        doc,
        [
            "Donner un niveau à tous les joueurs, même sans classement officiel.",
            "Traduire les classements nationaux vers une échelle mondiale commune.",
            "Utiliser les matchs joués pour rendre le rating de plus en plus fiable.",
            "Permettre le matchmaking, les challenges, les leaderboards et la progression personnelle.",
        ],
    )

    doc.add_heading("2. Pourquoi on ne peut pas copier les points officiels", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "Chaque pays utilise un système différent. En France, les points FFTT ne sont pas calculés comme le TTR allemand. Les points WTT/ITTF des pros ne sont pas un Elo : ce sont des points de ranking liés aux tournois, au calendrier, au rang atteint et aux points à défendre."
    )

    add_small_table(
        doc,
        ["Source", "Ce que ça représente", "Pourquoi ce n'est pas directement comparable"],
        [
            ["FFTT France", "Points nationaux et progression fédérale", "Barème français, dérive, compétitions locales."],
            ["TTR Allemagne", "Rating de force proche Elo", "Échelle allemande, population et compétitions propres."],
            ["WTT / ITTF", "Points de ranking professionnel", "Points liés aux tournois, pas à une probabilité directe de victoire."],
            ["Débutant", "Aucun classement officiel", "Le niveau doit être estimé puis calibré par matchs."],
        ],
        [1800, 3300, 4260],
    )

    doc.add_heading("3. Le modèle commun", level=1)
    p = doc.add_paragraph()
    p.add_run("Chaque joueur possède plusieurs informations distinctes. Le classement officiel reste visible, mais le PGR devient le score de référence dans l'app.")
    add_small_table(
        doc,
        ["Champ", "Rôle"],
        [
            ["Classement officiel", "Signal d'entrée : FFTT, TTR, WTT, RFETM, FITeT, etc."],
            ["PGR estimé", "Premier rating calculé à partir de la source officielle ou du questionnaire."],
            ["PGR actuel", "Rating mis à jour après les matchs Ping Pang et les matchs officiels importés."],
            ["Confiance / RD", "Niveau d'incertitude : fort au début, plus faible avec l'historique."],
            ["Source", "Permet de savoir si le rating vient d'un questionnaire, d'une fédération, du WTT ou de matchs Ping Pang."],
        ],
        [2600, 6760],
    )

    doc.add_heading("4. Fonctionnement général", level=1)
    add_numbered(
        doc,
        [
            "Le joueur s'inscrit et indique s'il possède un classement officiel.",
            "Si oui, l'app récupère ou vérifie la source officielle et calcule un PGR initial.",
            "Si non, l'app utilise un questionnaire de départ pour donner un PGR provisoire.",
            "Les premiers matchs servent de calibration : le score bouge plus vite au début.",
            "Plus le joueur accumule des matchs validés, plus le rating devient fiable.",
            "Le PGR sert ensuite au matchmaking, aux rankings, aux challenges et aux comparaisons.",
        ],
    )

    doc.add_heading("5. Pourquoi Glicko-2 plutôt qu'un Elo simple", level=1)
    p = doc.add_paragraph()
    p.add_run("Elo est simple et compréhensible, mais il gère mal l'incertitude. Glicko-2 ajoute une information essentielle : à quel point le rating est fiable.")
    add_small_table(
        doc,
        ["Concept", "Explication", "Effet dans l'app"],
        [
            ["Rating", "Le niveau estimé du joueur.", "Exemple : 1680 PGR."],
            ["RD", "L'incertitude autour du rating.", "Un nouveau joueur bouge vite, un joueur stable bouge moins."],
            ["Volatility", "La variabilité des performances.", "Utile pour détecter les joueurs irréguliers ou en forte progression."],
        ],
        [1800, 3600, 3960],
    )

    doc.add_heading("6. Cas des débutants", level=1)
    p = doc.add_paragraph()
    p.add_run("Un joueur sans classement officiel ne doit pas être exclu du système. Il reçoit un PGR provisoire, calculé à partir d'un questionnaire simple.")
    add_bullets(
        doc,
        [
            "Débutant complet : PGR bas, incertitude très forte.",
            "Joueur loisir régulier : PGR intermédiaire, incertitude forte.",
            "Joueur de club sans classement renseigné : PGR plus élevé, incertitude moyenne à forte.",
            "Après 5 à 10 matchs validés, le rating devient plus crédible.",
            "Après 20 à 30 matchs, le rating devient suffisamment stable pour le matchmaking.",
        ],
    )

    doc.add_heading("7. Fonctionnement précis pour les joueurs professionnels", level=1)
    add_callout(
        doc,
        "Règle fondamentale pour les pros",
        "Un joueur pro ne reçoit pas un PGR arbitraire parce qu'il est connu. Son PGR est calculé depuis ses sources officielles, puis recalibré avec ses résultats réels contre d'autres joueurs pros.",
    )

    doc.add_heading("7.1 Données utilisées pour un joueur pro", level=2)
    p = doc.add_paragraph()
    p.add_run("Pour les pros, la source principale est WTT/ITTF, car elle fournit le ranking mondial, les points et les résultats des compétitions internationales.")
    add_bullets(
        doc,
        [
            "Identifiant officiel WTT/ITTF si disponible.",
            "Nom, pays, genre/catégorie.",
            "Rang mondial à une date donnée.",
            "Points WTT/ITTF à une date donnée.",
            "Résultats de matchs en compétitions internationales.",
            "Adversaires rencontrés, score, date, compétition, tour, vainqueur.",
            "Historique mensuel du ranking si accessible.",
        ],
    )

    doc.add_heading("7.2 Pourquoi les points WTT ne deviennent pas directement le PGR", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Exemple : Sun Yingsha peut avoir 11 125 points WTT. Ces 11 125 points ne sont pas un Elo. Ils indiquent sa domination dans le circuit professionnel selon les règles WTT, mais ils ne disent pas directement : 'elle a 11 125 de niveau'."
    )
    p = doc.add_paragraph()
    p.add_run("Le PGR utilise ces points comme un signal, avec d'autres signaux : rang, historique, adversaires battus, régularité et récence des performances.")

    doc.add_heading("7.3 Calcul d'initialisation pour un pro", level=2)
    p = doc.add_paragraph()
    p.add_run("Le PGR initial d'un pro peut être calculé avec une fonction de scoring composée :")
    table = add_small_table(
        doc,
        ["Signal", "Rôle dans le calcul"],
        [
            ["Rang WTT/ITTF", "Signal le plus lisible : top 1, top 10, top 100, etc."],
            ["Points WTT/ITTF", "Signal de domination récente, souvent traité avec une fonction logarithmique."],
            ["Résultats récents", "Victoires/défaites contre des adversaires forts."],
            ["Qualité des adversaires", "Battre un top 10 vaut plus que battre un joueur beaucoup plus bas."],
            ["Récence", "Un résultat récent pèse plus qu'un résultat ancien."],
        ],
        [2500, 6860],
    )
    p = doc.add_paragraph()
    p.add_run("Formule conceptuelle :").bold = True
    p = doc.add_paragraph()
    r = p.add_run("PGR_initial = f(rang, log(points + 1), résultats récents, qualité des adversaires, récence)")
    r.font.name = "Courier New"
    r.font.size = Pt(9.5)

    doc.add_heading("7.4 Stabilisation avec les matchs officiels", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Après l'initialisation, on traite chronologiquement les matchs officiels WTT/ITTF avec Glicko-2. Le système ajuste progressivement le PGR pour que les probabilités de victoire correspondent aux résultats observés."
    )
    add_numbered(
        doc,
        [
            "On importe les rankings et les résultats WTT/ITTF sur une période donnée.",
            "On donne un rating initial à chaque pro.",
            "On rejoue les matchs dans l'ordre chronologique.",
            "Chaque victoire ou défaite met à jour le PGR et l'incertitude.",
            "Les joueurs actifs avec beaucoup de matchs obtiennent une confiance forte.",
        ],
    )

    doc.add_heading("7.5 Exemple concret : Sun Yingsha", level=2)
    add_small_table(
        doc,
        ["Étape", "Traitement Ping Pang"],
        [
            ["Source", "WTT/ITTF : rang mondial, points, historique de résultats."],
            ["Import", "Création ou mise à jour du profil joueur officiel."],
            ["Initialisation", "PGR calculé via la fonction WTT, pas choisi à la main."],
            ["Calibration", "Les matchs WTT/ITTF récents ajustent le rating."],
            ["Affichage app", "PGR, rang Ping Pang, rang WTT, source et niveau de confiance."],
        ],
        [2200, 7160],
    )
    p = doc.add_paragraph()
    p.add_run(
        "Ainsi, si Sun Yingsha est très haut dans le PGR, ce n'est pas parce que l'équipe l'a placée manuellement au sommet. C'est parce que les données officielles et ses résultats réels produisent mathématiquement un rating très élevé."
    )

    doc.add_heading("7.6 Cas d'un pro présent dans plusieurs sources", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Certains joueurs pros existent à la fois dans une source nationale et dans WTT/ITTF. Par exemple, un joueur français peut avoir un classement FFTT et un ranking international. Ping Pang doit fusionner ces identités dans un seul profil."
    )
    add_bullets(
        doc,
        [
            "WTT/ITTF sert d'ancre internationale.",
            "La source nationale enrichit l'historique local.",
            "Le matching se fait par identifiant officiel, nom normalisé, pays, club et validation manuelle si besoin.",
            "Le PGR final utilise les deux sources, mais pondère plus fortement les matchs de meilleure qualité et mieux reliés au graphe mondial.",
        ],
    )

    doc.add_heading("7.7 Pros, inactivité et baisse de classement officiel", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Un pro peut perdre des points WTT parce qu'il joue moins ou qu'il ne défend pas certains points de tournoi. Cela ne signifie pas forcément qu'il est devenu beaucoup moins fort. Le PGR doit donc séparer niveau estimé et incertitude."
    )
    add_bullets(
        doc,
        [
            "Si un pro devient inactif, son PGR ne s'effondre pas automatiquement.",
            "Son incertitude augmente : on est moins sûr de son niveau actuel.",
            "Quand il rejoue, ses nouveaux résultats recalibrent le rating.",
            "Cela évite de confondre absence du circuit et baisse réelle de niveau.",
        ],
    )

    doc.add_heading("7.8 Impact des pros sur le rating mondial", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Les pros sont essentiels parce qu'ils servent d'ancres mondiales. Les compétitions internationales relient les pays entre eux. Sans ces matchs ponts, chaque pays reste enfermé dans son propre système."
    )
    add_bullets(
        doc,
        [
            "Un joueur français qui affronte un joueur allemand crée un pont France-Allemagne.",
            "Un joueur européen qui affronte un joueur chinois en WTT crée un pont Europe-Asie.",
            "Ces ponts permettent au modèle d'apprendre les écarts entre systèmes nationaux.",
            "Plus il y a de matchs internationaux, plus le PGR devient stable mondialement.",
        ],
    )

    doc.add_heading("8. Fonctionnement deep : un pro chinois en concurrence avec un joueur français", level=1)
    add_callout(
        doc,
        "Principe de concurrence mondiale",
        "Un joueur chinois et un joueur français sont comparables uniquement parce qu'ils sont projetés sur la même variable cachée : leur force de jeu estimée. Le pays et la fédération ne donnent qu'un signal de départ ; les matchs créent les ponts qui rendent la comparaison stable.",
    )

    doc.add_heading("8.1 Le PGR comme force de jeu commune", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Le modèle ne dit pas : 'les points chinois valent plus que les points français'. Il dit : chaque joueur possède une force de jeu latente, appelée PGR, et chaque source officielle donne une estimation imparfaite de cette force."
    )
    p = doc.add_paragraph()
    r = p.add_run("force_reelle_du_joueur ≈ PGR")
    r.font.name = "Courier New"
    r.font.size = Pt(9.5)
    p = doc.add_paragraph()
    p.add_run(
        "Le rôle du système est de trouver le PGR qui explique le mieux les résultats observés : qui bat qui, à quelle date, dans quelle compétition, avec quel score et contre quel niveau d'adversaire."
    )

    doc.add_heading("8.2 Deux joueurs, deux sources différentes", level=2)
    add_small_table(
        doc,
        ["Joueur", "Source d'entrée", "Ce que l'app sait au départ", "Incertitude"],
        [
            [
                "Pro chinois",
                "WTT / ITTF",
                "Rang mondial, points, résultats internationaux, adversaires forts.",
                "Faible à moyenne si beaucoup de matchs récents.",
            ],
            [
                "Joueur français",
                "FFTT Smartping",
                "Points FFTT, club, historique de classement, matchs nationaux.",
                "Moyenne à forte si peu de matchs contre le graphe international.",
            ],
        ],
        [1700, 1900, 3900, 1860],
    )
    p = doc.add_paragraph()
    p.add_run(
        "Le pro chinois est souvent mieux connecté au graphe mondial parce qu'il joue des compétitions internationales. Le joueur français local peut être très bien mesuré dans le graphe français, mais moins bien relié au niveau mondial s'il ne joue jamais contre des étrangers ou des joueurs WTT."
    )

    doc.add_heading("8.3 La notion de graphe de matchs", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Le système construit un graphe : chaque joueur est un noeud, chaque match est une arête. Une victoire de A contre B crée une information directe sur l'écart de niveau entre A et B."
    )
    add_small_table(
        doc,
        ["Élément du graphe", "Exemple", "Utilité"],
        [
            ["Noeud", "Sun Yingsha, Alexis Lebrun, joueur FFTT 1800.", "Représente un joueur."],
            ["Arête", "A bat B 3-1 en WTT, ou C bat D en championnat FFTT.", "Mesure un écart de force."],
            ["Poids", "Compétition officielle récente > match ancien incomplet.", "Contrôle la fiabilité de l'information."],
            ["Pont", "Un Français joue en WTT contre un Chinois.", "Relie deux écosystèmes nationaux."],
        ],
        [2100, 3300, 3960],
    )

    doc.add_heading("8.4 Les ponts France-Chine", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Pour comparer un joueur chinois et un joueur français, il faut des ponts. Les meilleurs ponts sont les matchs internationaux, les joueurs qui apparaissent dans plusieurs systèmes et les championnats où des joueurs étrangers participent."
    )
    add_bullets(
        doc,
        [
            "WTT / ITTF : les meilleurs joueurs chinois affrontent des joueurs français, allemands, japonais, coréens, etc.",
            "Joueurs français présents en WTT : ils relient la FFTT au graphe mondial.",
            "Joueurs étrangers en championnat français : ils relient leur niveau international au niveau FFTT.",
            "Matchs Ping Pang internationaux : plus l'app grandit, plus elle crée ses propres ponts.",
        ],
    )

    doc.add_heading("8.5 Apprendre la conversion au lieu de l'inventer", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Au départ, chaque source a une fonction de traduction vers le PGR. Mais cette fonction ne doit pas rester fixée au feeling. Elle doit être ajustée avec les résultats réels."
    )
    p = doc.add_paragraph()
    r = p.add_run("PGR_initial = a_source + b_source × transformation(classement_source)")
    r.font.name = "Courier New"
    r.font.size = Pt(9.5)
    p = doc.add_paragraph()
    p.add_run("Exemples de transformations :")
    add_bullets(
        doc,
        [
            "Pour un rating proche Elo comme TTR : transformation plutôt linéaire.",
            "Pour des points WTT : transformation logarithmique, car 11 000 points ne valent pas 11 fois 1 000 points.",
            "Pour un rang : transformation par percentile ou log du rang.",
            "Pour la FFTT : transformation par courbe ou tranches, puis ajustement avec les matchs.",
        ],
    )
    p = doc.add_paragraph()
    p.add_run(
        "Les paramètres a_source et b_source sont ensuite corrigés pour minimiser les erreurs de prédiction : si le modèle sous-estime systématiquement les joueurs français face aux joueurs internationaux, la conversion FFTT est ajustée."
    )

    doc.add_heading("8.6 Exemple concret : pro chinois vs joueur français", level=2)
    add_small_table(
        doc,
        ["Étape", "Pro chinois", "Joueur français"],
        [
            [
                "Import",
                "Profil WTT/ITTF, rang mondial, résultats internationaux.",
                "Profil FFTT, points nationaux, historique de matchs français.",
            ],
            [
                "PGR initial",
                "Calculé depuis rang WTT, points WTT et résultats récents.",
                "Calculé depuis points FFTT et historique national.",
            ],
            [
                "Incertitude",
                "Faible si beaucoup de matchs WTT récents.",
                "Plus élevée si peu de ponts avec le niveau international.",
            ],
            [
                "Calibration",
                "Matchs contre des joueurs de plusieurs pays.",
                "Matchs FFTT + éventuels matchs contre joueurs reliés au graphe WTT.",
            ],
            [
                "Classement app",
                "Trié par PGR sur le leaderboard mondial.",
                "Trié par PGR sur la même échelle, avec badge de confiance.",
            ],
        ],
        [1700, 3830, 3830],
    )

    doc.add_heading("8.7 Si les deux joueurs s'affrontent dans Ping Pang", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Quand le pro chinois et le joueur français jouent un match validé dans l'app, le système calcule d'abord la probabilité attendue de victoire à partir de leurs PGR."
    )
    p = doc.add_paragraph()
    r = p.add_run("P(victoire chinois) = 1 / (1 + 10^((PGR_français - PGR_chinois) / 400))")
    r.font.name = "Courier New"
    r.font.size = Pt(9.5)
    add_small_table(
        doc,
        ["Situation", "Effet sur le PGR"],
        [
            [
                "Le pro chinois gagne comme prévu",
                "Petit gain ou quasi stabilité : le résultat confirme l'attendu.",
            ],
            [
                "Le Français perd mais fait un match serré",
                "Si le score détaillé est utilisé, légère correction possible, sinon effet limité.",
            ],
            [
                "Le Français gagne",
                "Gros upset : le Français monte fortement, le pro chinois baisse, surtout si le match est fiable.",
            ],
            [
                "Le Français a peu de matchs",
                "Son RD est élevé : son PGR peut bouger beaucoup après le match.",
            ],
            [
                "Le pro chinois a beaucoup de matchs",
                "Son RD est faible : son PGR bouge moins violemment, sauf upset répété.",
            ],
        ],
        [3100, 6260],
    )

    doc.add_heading("8.8 Pourquoi ce système est stable", level=2)
    add_bullets(
        doc,
        [
            "Le PGR ne dépend pas d'une seule source : il combine ranking, résultats, adversaires et récence.",
            "Les pros internationaux servent d'ancres parce qu'ils affrontent des joueurs de nombreux pays.",
            "Les joueurs nationaux sont reliés progressivement via les ponts : WTT, championnats, étrangers en club, matchs Ping Pang.",
            "L'incertitude évite de surclasser ou sous-classer trop vite un joueur peu connecté.",
            "Le modèle se corrige avec les résultats : si une conversion pays est mauvaise, les matchs la révèlent.",
        ],
    )

    doc.add_heading("8.9 Affichage produit dans l'app", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Pour que le système soit compréhensible par l'utilisateur, il faut afficher le PGR avec son niveau de confiance."
    )
    add_small_table(
        doc,
        ["Élément affiché", "Exemple"],
        [
            ["PGR", "2865"],
            ["Statut", "Confirmé / Provisoire / En calibration"],
            ["Source principale", "WTT/ITTF, FFTT, Matchs Ping Pang"],
            ["Confiance", "Élevée, moyenne, faible"],
            ["Classements", "Monde, pays, ville, club, catégorie"],
        ],
        [2700, 6660],
    )

    doc.add_heading("9. Sources data à construire en premier", level=1)
    add_small_table(
        doc,
        ["Priorité", "Source", "Pourquoi"],
        [
            ["1", "WTT / ITTF", "Ancre mondiale et données pros."],
            ["1", "FFTT Smartping", "Données françaises riches : joueurs, classements, matchs."],
            ["2", "Allemagne TTR", "Très bon signal de force, proche Elo."],
            ["2", "Angleterre / Rankedin", "Données structurées et utiles pour plusieurs pays."],
            ["3", "Italie / Espagne", "Complétion Europe."],
            ["4", "Chine / Japon / Corée", "WTT au départ, partenariats ensuite."],
        ],
        [1200, 2600, 5560],
    )

    doc.add_heading("10. Version à construire pour prouver le concept", level=1)
    add_bullets(
        doc,
        [
            "Importer un petit échantillon WTT/ITTF de joueurs pros.",
            "Importer ou simuler des joueurs FFTT français.",
            "Créer des PGR initiaux selon la source.",
            "Simuler ou importer des matchs historiques.",
            "Afficher PGR, niveau de confiance, source et historique d'évolution.",
            "Ajouter un parcours débutant sans classement officiel.",
        ],
    )

    doc.add_heading("11. Phrase de synthèse", level=1)
    add_callout(
        doc,
        "Positionnement",
        "Le PGR n'est pas un copier-coller des classements existants. C'est un rating global indépendant, relié aux sources officielles, ajusté par les résultats réels, et capable d'intégrer aussi bien un débutant qu'un joueur professionnel.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
